#!/usr/bin/env python3
"""
Build editable .docx files for the 1-page teaser résumés + the IA cover letter.

Outputs to:
  C:\\Users\\riket\\OneDrive\\Desktop\\Organized\\5 - Personal Development\\01 - Resume & Cover Letters\\Generated\\

Files generated:
  Riket B Patel — Résumé.docx
  Riket B Patel — Résumé (ADP, Job 597814).docx
  Riket B Patel — Résumé (Internet Archive, Program Coordinator).docx
  Riket B Patel — Cover Letter (Internet Archive, Program Coordinator).docx

Run:
  python scripts/build_resume_docx.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = Path(
    r"C:\Users\riket\OneDrive\Desktop\Organized\5 - Personal Development\01 - Resume & Cover Letters\Generated"
)
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# ─── Style helpers ─────────────────────────────────────────────────────────

INK = RGBColor(0x11, 0x11, 0x11)
MUTED = RGBColor(0x55, 0x55, 0x55)
ACCENT = RGBColor(0x1A, 0x3A, 0x2A)
HIGHLIGHT = RGBColor(0xB8, 0x86, 0x0B)
RULE = RGBColor(0x1A, 0x1A, 0x1A)

FONT = "Calibri"


def set_run(run, *, size=10.5, bold=False, italic=False, color=INK, font=FONT):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def add_para(doc, *, space_before=0, space_after=2, line_spacing=1.15, align=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if align is not None:
        p.alignment = align
    return p


def add_hr(p):
    """Add a horizontal rule below the paragraph (bottom border)."""
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")  # 1pt
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1A1A1A")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_hr_top(p):
    """Add a horizontal rule above the paragraph."""
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "12")  # 1.5pt
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), "1A1A1A")
    pBdr.append(top)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1A1A1A")
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_margins(doc, top=0.5, bottom=0.5, left=0.55, right=0.55):
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)


def section_header(doc, text):
    p = add_para(doc, space_before=8, space_after=2)
    r = p.add_run(text.upper())
    set_run(r, size=10, bold=True, color=INK)
    r.font.color.rgb = INK
    # Letter spacing via XML
    rPr = r._r.get_or_add_rPr()
    spc = OxmlElement("w:spacing")
    spc.set(qn("w:val"), "28")  # ~1.4pt letter spacing
    rPr.append(spc)
    add_hr(p)
    return p


def name_header(doc, name):
    p = add_para(doc, space_before=0, space_after=2, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.0)
    r = p.add_run(name.upper())
    set_run(r, size=22, bold=True, color=INK)
    # Wide letter spacing
    rPr = r._r.get_or_add_rPr()
    spc = OxmlElement("w:spacing")
    spc.set(qn("w:val"), "40")
    rPr.append(spc)


def contact_line(doc, parts):
    p = add_para(doc, space_after=2, line_spacing=1.2)
    r = p.add_run(" · ".join(parts))
    set_run(r, size=9.5, color=MUTED)


def headline_bar(doc, text):
    p = add_para(doc, space_before=4, space_after=4, line_spacing=1.2, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run(text.upper())
    set_run(r, size=11, bold=True, color=ACCENT)
    rPr = r._r.get_or_add_rPr()
    spc = OxmlElement("w:spacing")
    spc.set(qn("w:val"), "16")
    rPr.append(spc)
    add_hr_top(p)


def summary_para(doc, text_runs):
    """text_runs: list of (text, bold_bool) tuples."""
    p = add_para(doc, space_before=4, space_after=2, line_spacing=1.25)
    for text, bold in text_runs:
        r = p.add_run(text)
        set_run(r, size=10.5, bold=bold, color=INK)


def bullet_with_quote(doc, lead_runs, quote_text=None):
    """Bullet starting with a triangle, optional italic quote at end."""
    p = add_para(doc, space_before=1, space_after=1, line_spacing=1.25)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.18)

    arrow = p.add_run("▸  ")
    set_run(arrow, size=10.5, color=HIGHLIGHT, bold=True)

    for text, bold in lead_runs:
        r = p.add_run(text)
        set_run(r, size=10.5, bold=bold, color=INK)

    if quote_text:
        # leading space + italic quote
        q = p.add_run(" " + quote_text)
        set_run(q, size=10.5, italic=True, color=RGBColor(0x2A, 0x2A, 0x2A))


def job_block(doc, role, dates, body_runs):
    """role + right-aligned dates on same line via tab, then body paragraph."""
    p = add_para(doc, space_before=4, space_after=0, line_spacing=1.2)
    # Right tab stop near right margin
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(7.4), WD_ALIGN_PARAGRAPH.RIGHT)

    role_run = p.add_run(role)
    set_run(role_run, size=10.5, bold=True, color=INK)
    p.add_run("\t")
    date_run = p.add_run(dates)
    set_run(date_run, size=9.5, italic=True, color=MUTED)

    body = add_para(doc, space_before=1, space_after=2, line_spacing=1.25)
    for text, bold in body_runs:
        r = body.add_run(text)
        set_run(r, size=10.5, bold=bold, color=INK)


def two_column_block(doc, left_lines, right_lines):
    """Two-column layout via a borderless table."""
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    # Set column widths
    table.columns[0].width = Inches(3.65)
    table.columns[1].width = Inches(3.65)
    cells = table.rows[0].cells
    cells[0].width = Inches(3.65)
    cells[1].width = Inches(3.65)

    def fill(cell, lines):
        cell.text = ""  # clear default
        for i, line in enumerate(lines):
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.25
            for text, bold in line:
                r = p.add_run(text)
                set_run(r, size=10.5, bold=bold, color=INK)

    fill(cells[0], left_lines)
    fill(cells[1], right_lines)

    # Remove table borders
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def footnote(doc, text):
    p = add_para(doc, space_before=6, space_after=0, line_spacing=1.2, align=WD_ALIGN_PARAGRAPH.RIGHT)
    r = p.add_run(text)
    set_run(r, size=9, italic=True, color=MUTED)


# ─── Shared header builder ─────────────────────────────────────────────────

CONTACT_PARTS = [
    "56 Benford Ln, Edgewater Park, NJ 08010",
    "(267) 408-6295",
    "Riketpatel@gmail.com",
    "linkedin.com/in/riketpatel",
    "riketpatel.com",
    "Trilingual: English, Hindi, Gujarati",
]


def build_header(doc, headline):
    name_header(doc, "Riket B. Patel")
    contact_line(doc, CONTACT_PARTS)
    headline_bar(doc, headline)


# ─── Shared education + certifications + footnote ──────────────────────────

def edu_certs_block(doc):
    section_header(doc, "Education & Certifications")
    left = [
        [("Pennsylvania State University", True)],
        [("M.P.S. Data Analytics — May 2018", False)],
        [("Graduate Certificate, Business Analytics — Aug 2017", False)],
        [("B.S. Business Administration (Mgmt. & Marketing) — May 2014", False)],
    ]
    right = [
        [("Certified AI Product Management — Dec 2024", True)],
        [("Certified Agile Leader (CAL-1), Scrum Alliance — Jun 2024", True)],
        [("Mental Health First Aider — May 2024", True)],
        [("AWS Cloud Practitioner Path · NJ Real Estate Salesperson", False)],
    ]
    two_column_block(doc, left, right)


def references_footnote(doc, extra=""):
    text = "References available on request."
    if extra:
        text += " " + extra
    footnote(doc, text)


# ─── BUILDER 1: Main résumé ────────────────────────────────────────────────

def build_main_resume():
    doc = Document()
    set_margins(doc)

    build_header(doc, "Product Manager  ·  Healthcare Tech  ·  Multi-Venture Operator")

    summary_para(doc, [
        ("Spent seven years at Merck building clinical-data systems for teams in six countries — work that eliminated about ", False),
        ("10,000 hours", True),
        (" of annual manual effort and cut a China NDA cycle from ", False),
        ("6-7 months to about one month", True),
        (". Left Merck in November 2025 to run ", False),
        ("Legacy Bridge Alliance Group", True),
        (" full time: a New Jersey real-estate practice (HomeSmart affiliate), a business-brokerage practice, two consumer Shopify stores, and four free public-good sites. First-generation Indian-American, ADHD-aware operator. M.S. Data Analytics, Penn State; Certified AI Product Manager + Agile Leader (CAL-1). Open to product or program roles where being trusted with the mission matters more than chasing a title.", False),
    ])

    section_header(doc, "Selected Wins")

    bullet_with_quote(doc,
        [("Scale IT for Clinical Portfolio", True),
         (" — eliminated ~10,000 hours of annual manual effort across Delta Fusion, CDI, CDDR, and SLS; delivered ahead of plan; projected to double to 20,000 hours. Senior leadership endorsement (Jill McGinn, Director of Corporate Strategy, Merck):", False)],
        '"Strategic thinker, unbound by convention; values diversity and inclusion; goal-oriented and ambitious."'
    )
    bullet_with_quote(doc,
        [("A&R Submission Translation (ARST)", True),
         (" — cut China NDA translation cycle from ", False),
         ("6-7 months to about one month", True),
         (". User research with stakeholders in Tokyo, Shanghai, Berlin, Brussels, and Boston.", False)]
    )
    bullet_with_quote(doc,
        [("Squad Lead (2022)", True),
         (" — owned sprint planning, delivery cadence, and mentorship for the cross-functional product squad. Year objectives shipped on schedule with zero audit findings across 3 inspection-sensitive systems. 20+ peer recognitions on Merck's INSPIRE platform — and an even higher count given to others.", False)]
    )
    bullet_with_quote(doc,
        [("Reverse-mentor feedback", True),
         (" from Robert Wiley, Director of Oncology National Account at Merck:", False)],
        '"exceptional in every area we explored…I do not have any further development feedback."'
    )
    bullet_with_quote(doc,
        [("Summer-intern mentor", True),
         (" to Jenny Mao (Cornell Tech master’s student, 2025):", False)],
        '"I truly don\'t think I would have completed my internship so successfully without your mentorship."'
    )

    section_header(doc, "Experience")

    job_block(doc,
        "Product Manager, CRWEG — Clinical Trial Analysis & Reporting · Merck & Co., Inc.",
        "Aug 2018 – Nov 2025",
        [("Owned the product roadmap for 5 clinical-data systems serving Biostatistics, Medical Writing, and Research Decision Sciences across 6 countries. Promoted to ", False),
         ("Squad Lead", True),
         (" (2022). ", False),
         ("SME", True),
         (" for BARDS regulatory audit-readiness across 3 inspection-sensitive systems. ", False),
         ("TCO Data Steward", True),
         (" for Value Team using ", False),
         ("Apptio + Ariba", True),
         (". Change Champion for O365 / Microsoft Teams enterprise adoption. Represented the product line at international strategy sessions in ", False),
         ("Prague, Brussels, Zurich", True),
         (".", False)]
    )

    job_block(doc,
        "Solution Specialist II → EHR Application Specialist · NextGen Healthcare, Horsham, PA",
        "May 2014 – Apr 2017",
        [("Mirth Connect & EHR client implementations. HL7 / FHIR / SQL Server troubleshooting. Built deep domain expertise in ", False),
         ("HIPAA, ICD-10, Health Information Exchange, and clinical documentation", True),
         (" — foundation for the later regulated-environment work.", False)]
    )

    job_block(doc,
        "Founder · Legacy Bridge Alliance Group",
        "2025 – present",
        [("Left Merck in November 2025 to take on full-time what was already taking up nights and weekends. Now run ", False),
         ("Metta Realty Partners", True),
         (" (NJ real estate, HomeSmart affiliate) and ", False),
         ("Metta Legacy Partners", True),
         (" (business brokerage that handles the operating company while Metta Realty handles the property when both sell together), two consumer Shopify stores (KiddieSketch, KiddieGo), and four free public-good sites: KiddieWordle (17 languages), BestPythonCourse, LoanRatesFinder, FreeRateFinder. Also setting up a Python and generative-AI teacher position in my family's village in Anand, Gujarat — continuing my father's promise to give back there.", False)]
    )

    edu_certs_block(doc)

    references_footnote(doc, "Longer narrative, side projects, and stories at riketpatel.com.")

    out = OUTPUT_DIR / "Riket B Patel — Résumé.docx"
    doc.save(out)
    print(f"✓ {out.name}")


# ─── BUILDER 2: ADP-tailored résumé ────────────────────────────────────────

def build_adp_resume():
    doc = Document()
    set_margins(doc)

    build_header(doc, "Product Manager  ·  Enterprise SaaS  ·  Regulated Data & AI")

    summary_para(doc, [
        ("Seven years owning enterprise SaaS product roadmaps at Merck across teams in six countries. Removed ~", False),
        ("10,000 hours", True),
        (" of annual manual effort and cut a regulated submission cycle from ", False),
        ("6-7 months to about one month", True),
        (". Familiar with SOX-adjacent compliance, multi-jurisdiction localization, audit-readiness, and AI / ML integration. Left Merck in November 2025 to run a small holding company of real-estate, brokerage, and e-commerce ventures — comfortable being the central operations person who keeps everything on schedule. First-generation Indian-American, ADHD-aware operator. M.S. Data Analytics, Penn State. Certified AI Product Manager + Agile Leader (CAL-1).", False),
    ])

    section_header(doc, "Selected Wins (HCM SaaS-Adjacent)")

    bullet_with_quote(doc,
        [("Scale IT for Clinical Portfolio", True),
         (" — eliminated ~10,000 hours of annual manual effort across 4 interconnected systems; same operating pattern as automating recurring payroll-prep, time-attestation, and reporting workflows at scale. Senior leadership endorsement (Jill McGinn, Director of Corporate Strategy, Merck):", False)],
        '"Strategic thinker, unbound by convention; values diversity and inclusion; goal-oriented and ambitious."'
    )
    bullet_with_quote(doc,
        [("A&R Submission Translation", True),
         (" — cut China NDA translation cycle from ", False),
         ("6-7 months to about one month", True),
         ("; HCM analog: multi-jurisdiction localization for payroll filings, benefits docs, onboarding flows.", False)]
    )
    bullet_with_quote(doc,
        [("TCO Data Steward", True),
         (" for Value Team using ", False),
         ("Apptio + Ariba", True),
         (" — aligned financial models with executive reporting. Direct fit for ADP enterprise SaaS financial modeling work.", False)]
    )
    bullet_with_quote(doc,
        [("BARDS audit-readiness SME", True),
         (" across 3 inspection-sensitive systems (audit trail, user access, data integrity) — comparable rigor to ADP's SOC 2 / IRS / state-payroll audit posture.", False)]
    )
    bullet_with_quote(doc,
        [("Squad Lead (2022)", True),
         (" — managed sprint cadence and delivery accountability for a cross-functional team. Year objectives shipped on schedule; zero audit findings across 3 inspection-sensitive systems; 20+ peer recognitions. Reverse-mentee (Robert Wiley, Director, Merck):", False)],
        '"exceptional in every area we explored."'
    )

    section_header(doc, "Experience")

    job_block(doc,
        "Product Manager, CRWEG · Merck & Co., Inc., Upper Gwynedd, PA",
        "Aug 2018 – Nov 2025",
        [("Enterprise SaaS product manager across 5 interconnected clinical-data systems serving 15+ stakeholder groups in 6 countries. Promoted to ", False),
         ("Squad Lead", True),
         (" (2022). Change Champion for O365 / Microsoft Teams enterprise adoption. International strategy sessions in ", False),
         ("Prague, Brussels, Zurich", True),
         (" on cloud modernization and AI / ML integration.", False)]
    )

    job_block(doc,
        "Solution Specialist II → EHR Application Specialist · NextGen Healthcare, Horsham, PA",
        "May 2014 – Apr 2017",
        [("Mirth Connect & EHR client implementations. HL7 / FHIR / SQL Server troubleshooting. Domain expertise in ", False),
         ("HIPAA, ICD-10, Health Information Exchange, clinical documentation", True),
         (".", False)]
    )

    job_block(doc,
        "Founder · Legacy Bridge Alliance Group",
        "2025 – present",
        [("Left Merck in November 2025. Operate three small-business ventures simultaneously: ", False),
         ("Metta Realty Partners", True),
         (" (NJ real estate, HomeSmart affiliate), ", False),
         ("Metta Legacy Partners", True),
         (" (business brokerage), and two consumer Shopify stores (KiddieSketch, KiddieGo). Real-world experience with payroll, benefits, multi-state filings, and small-business compliance — the daily reality of ADP's customer base. Also run four free public-good content sites in evenings.", False)]
    )

    edu_certs_block(doc)
    references_footnote(doc, "Longer narrative and side projects at riketpatel.com.")

    out = OUTPUT_DIR / "Riket B Patel — Résumé (ADP, Job 597814).docx"
    doc.save(out)
    print(f"✓ {out.name}")


# ─── BUILDER 3: Internet Archive résumé ────────────────────────────────────

def build_ia_resume():
    doc = Document()
    set_margins(doc)

    build_header(doc, "Program Coordination  ·  Cross-Functional Operations  ·  Free-Access Builder")

    summary_para(doc, [
        ("Seven years coordinating multi-system product programs at Merck across six countries. Removed ~", False),
        ("10,000 hours", True),
        (" of annual manual effort and cut a regulated submission cycle from ", False),
        ("6-7 months to about one month", True),
        (". Left Merck in November 2025 to run small ventures full-time and build free things on the side. Currently setting up a Python and generative-AI teacher position in my family's village in Anand, Gujarat — continuing my father's promise to give back there. First-generation Indian-American, ADHD-aware operator, trilingual. Universal access to knowledge isn't a pivot — it's already the lane.", False),
    ])

    section_header(doc, "Selected Wins (Program-Ops Translatable)")

    bullet_with_quote(doc,
        [("Scale IT for Clinical Portfolio", True),
         (" — coordinated a 4-system program that eliminated ~10,000 hours of annual manual effort; same operating pattern as freeing a small staff to do mission work instead of repetitive prep. Senior leadership endorsement (Jill McGinn, Director of Corporate Strategy, Merck):", False)],
        '"Strategic thinker, unbound by convention; values diversity and inclusion; goal-oriented and ambitious."'
    )
    bullet_with_quote(doc,
        [("A&R Submission Translation", True),
         (" — cut a regulated translation cycle from ", False),
         ("6-7 months to about one month", True),
         (". User research with stakeholders in Tokyo, Shanghai, Berlin, Brussels, and Boston. Comfortable building for users whose language, time zone, and context are not my own.", False)]
    )
    bullet_with_quote(doc,
        [("BARDS audit-readiness SME", True),
         (" across 3 systems (audit trail, user access, data integrity) — same documentation rigor a grant report or accession workflow requires.", False)]
    )
    bullet_with_quote(doc,
        [("Squad Lead (2022)", True),
         (" — coordinated delivery and mentorship for a cross-functional team. Year objectives shipped on schedule with zero audit findings. Mentored Cornell Tech master's intern Jenny Mao through her 2025 Merck summer:", False)],
        '"I truly don\'t think I would have completed my internship so successfully without your mentorship."'
    )
    bullet_with_quote(doc,
        [("Free public-good sites", True),
         (": KiddieWordle (Wordle for kids in 17 languages, no in-play ads), BestPythonCourse, LoanRatesFinder, FreeRateFinder. Same instinct the Archive operates on, on a smaller scale.", False)]
    )

    section_header(doc, "Experience")

    job_block(doc,
        "Program Lead / Product Manager, CRWEG · Merck & Co., Inc., Upper Gwynedd, PA",
        "Aug 2018 – Nov 2025",
        [("Coordinated a multi-system product program for Biostatistics, Medical Writing, and Research Decision Sciences in 6 countries. Promoted to ", False),
         ("Squad Lead", True),
         (" (2022). International working sessions in Prague, Brussels, Zurich. Active member of Interfaith Organization, Next Gen Network, and capABILITY Network (disability inclusion) ERGs.", False)]
    )

    job_block(doc,
        "Solution Specialist II → EHR Application Specialist · NextGen Healthcare, Horsham, PA",
        "May 2014 – Apr 2017",
        [("Healthcare-client coordination, Mirth Connect / EHR implementations, HL7 / FHIR / SQL. Domain knowledge in HIPAA, ICD-10, Health Information Exchange, clinical documentation. Authored SOPs aligned to Public Health policy.", False)]
    )

    job_block(doc,
        "Founder · Legacy Bridge Alliance Group + free public-good sites",
        "2025 – present",
        [("One-person program-ops practice running on JIRA / Make.com / Google Workspace. Coordinate vendors, partners, fulfillment, and editorial across multiple workstreams. Currently also setting up a Python & generative-AI teacher position in my family's village in Anand, Gujarat — continuing my father's promise to give back there.", False)]
    )

    edu_certs_block(doc)
    references_footnote(doc, "Cover letter pairs with this résumé. Longer narrative at riketpatel.com.")

    out = OUTPUT_DIR / "Riket B Patel — Résumé (Internet Archive, Program Coordinator).docx"
    doc.save(out)
    print(f"✓ {out.name}")


# ─── BUILDER 4: Internet Archive cover letter ──────────────────────────────

def build_ia_cover_letter():
    doc = Document()
    set_margins(doc, top=0.75, bottom=0.75, left=0.85, right=0.85)

    # Sender header
    p = add_para(doc, space_after=2, line_spacing=1.0)
    r = p.add_run("Riket B. Patel")
    set_run(r, size=19, bold=True, color=INK)
    rPr = r._r.get_or_add_rPr()
    spc = OxmlElement("w:spacing")
    spc.set(qn("w:val"), "24")
    rPr.append(spc)

    contact_line(doc, [
        "56 Benford Ln, Edgewater Park, NJ 08010",
        "(267) 408-6295",
        "Riketpatel@gmail.com",
        "linkedin.com/in/riketpatel",
        "riketpatel.com",
    ])

    # Date
    p = add_para(doc, space_before=16, space_after=12, line_spacing=1.3)
    r = p.add_run("May 14, 2026")
    set_run(r, size=11, color=INK)

    # To block
    for line in [
        "Hiring Team",
        "Internet Archive",
        "300 Funston Avenue",
        "San Francisco, CA 94118",
    ]:
        p = add_para(doc, space_after=0, line_spacing=1.3)
        r = p.add_run(line)
        set_run(r, size=11, color=INK)

    # Salutation
    p = add_para(doc, space_before=16, space_after=8, line_spacing=1.4)
    r = p.add_run("Dear Hiring Team,")
    set_run(r, size=11.5, color=INK)

    # Paragraphs
    def body_para(text_runs):
        p = add_para(doc, space_after=8, line_spacing=1.5)
        for text, bold, italic in text_runs:
            r = p.add_run(text)
            set_run(r, size=11.5, bold=bold, italic=italic, color=INK)

    def subhead(text):
        p = add_para(doc, space_before=10, space_after=4, line_spacing=1.4)
        r = p.add_run(text)
        set_run(r, size=12, bold=True, color=ACCENT)

    body_para([
        ("I am writing about the Program Coordinator role at the Internet Archive. I spent seven years at Merck running multi-system product programs that eliminated about 10,000 hours of manual effort a year across teams in the US, Europe, and Asia-Pacific. I left Merck in November 2025 to run a small holding company of ventures full-time while figuring out the next chapter, and the work I already do on my own time — quietly, on nights and weekends — points in exactly the same direction the Archive does. That's why this role feels less like a career pivot and more like a continuation.", False, False),
    ])

    subhead("My professional journey")
    body_para([
        ("At Merck & Co., I led the Scale IT for Clinical Portfolio program — coordinating delivery across four interconnected clinical-data systems (Delta Fusion, CDI, CDDR, SLS) and 15+ stakeholder groups across six countries. The program eliminated about 10,000 hours of annual manual effort, was delivered ahead of plan, and is on track to double to 20,000 hours. I also designed and shipped the A&R Submission Translation platform, which cut translated Clinical Study Reports and supporting NDA documents from about 6-7 months to about one month — accelerating regulatory filings to China's NMPA, and the medicines those filings unlock. The research behind it was hands-on stakeholder work in Tokyo, Shanghai, Berlin, Brussels, and Boston. I was promoted to Squad Lead in 2022, served as the regulatory subject-matter expert for audit-readiness across three inspection-sensitive systems, and consistently shipped year objectives ahead of schedule with zero audit findings. My summer intern Jenny Mao — a Cornell Tech master's student spending the summer of 2025 on our team — later wrote: ", False, False),
        ('"I truly don\'t think I would have completed my internship so successfully without your mentorship."', False, True),
        (" That sentence is the one I'm proudest of from the whole tenure.", False, False),
    ])

    subhead("Why the Internet Archive")
    body_para([
        ("I'm a first-generation Indian-American. My family came to the US from a small village in Anand, Gujarat. At his farewell event there, my father made a public promise to always give back to that village. I am continuing that promise: I am setting up a Python and generative-AI teacher position in our village so the kids growing up where my parents grew up have access to the same tools the rest of the world is learning. That is the same instinct the Internet Archive operates on, just at a different scale — useful tools and knowledge should be free, reachable, and built for the people who need them most.", False, False),
    ])

    body_para([
        ("On my own time I also run a small set of free, ad-light public-good projects: ", False, False),
        ("KiddieWordle", True, False),
        (", a Wordle for younger players in seventeen languages with no ads in the play area; ", False, False),
        ("BestPythonCourse", True, False),
        (", plain-English guides for people learning Python without being upsold into a $300 course; and ", False, False),
        ("LoanRatesFinder", True, False),
        (" and ", False, False),
        ("FreeRateFinder", True, False),
        (", free mortgage- and loan-rate references with simple calculators. None of those will make anyone rich. That's the point. I keep building them anyway.", False, False),
    ])

    body_para([
        ("Longer term, I want to start a program that connects pharma — the industry where I grew professionally — with the Vedic community I grew up in. Two worlds that don't talk to each other often enough; both have things to learn from the other about discipline, ethics, and what we owe to the people we serve. I am not sure yet whether that lives as a nonprofit, an industry group, or just a recurring conversation, but the instinct behind it is the same one that drew me to the Archive: useful knowledge belongs to more people than currently have access to it.", False, False),
    ])

    subhead("Personal motivation and authenticity")
    body_para([
        ("As someone who navigates life with ADHD, I bring a unique perspective to product and program design — I naturally think about cognitive load, accessibility, and user experience through a lens of lived experience. This drives my commitment to building products and programs that work for everyone, which is the same commitment that drew me to the Internet Archive's work. I am also a Certified Mental Health First Aider (2024); I picked that up because too many people I respected were struggling silently and I wanted to be useful before they had to ask.", False, False),
    ])

    subhead("A note on the title")
    body_para([
        ("My most recent title at Merck was Product Manager, not Program Coordinator. I'd rather put that in front of you honestly than dress it up. The skills underneath — multi-stakeholder coordination, documentation discipline, working in regulated and audit-sensitive contexts, mentoring — are the ones this role seems to ask for. If the role is more senior than the title suggests, even better. If it's more entry-level than I'd guess from the description, I am also willing to have that conversation honestly. I'd rather start at the right level at an organization whose mission I'd be proud to point my son toward someday than chase a title at one I wouldn't.", False, False),
    ])

    body_para([
        ("My résumé, tailored to this role, is at riketpatel.com/resume/internet-archive/. My broader portfolio of side projects, including the free sites mentioned above, is at riketpatel.com. I would welcome a conversation at your convenience and look forward to hearing from you.", False, False),
    ])

    body_para([
        ("Thank you for the time and the work you do.", False, False),
    ])

    # Signoff
    p = add_para(doc, space_before=10, space_after=4, line_spacing=1.4)
    r = p.add_run("Best regards,")
    set_run(r, size=11.5, color=INK)

    p = add_para(doc, space_after=2, line_spacing=1.3)
    r = p.add_run("Riket B. Patel")
    set_run(r, size=11.5, bold=True, color=INK)

    # Personal sign-off
    p = add_para(doc, space_before=4, space_after=4, line_spacing=1.3)
    r = p.add_run("Hari Om Tat Sat.")
    set_run(r, size=10.5, italic=True, color=MUTED)

    for line in [
        "Product Manager · Master's in Data Analytics, Pennsylvania State University",
        "Certified Artificial Intelligence Product Manager · Certified Agile Leader (CAL-1), Scrum Alliance",
        "Certified Mental Health First Aider · NJ Real Estate Salesperson",
    ]:
        p = add_para(doc, space_after=0, line_spacing=1.3)
        r = p.add_run(line)
        set_run(r, size=10.5, color=MUTED)

    out = OUTPUT_DIR / "Riket B Patel — Cover Letter (Internet Archive, Program Coordinator).docx"
    doc.save(out)
    print(f"✓ {out.name}")


# ─── Main ──────────────────────────────────────────────────────────────────

# ─── BUILDER 5: NJ Judiciary BA / IT Analyst 2 résumé ─────────────────────

def build_nj_judiciary_resume():
    doc = Document()
    set_margins(doc)

    # Custom header so we can include US Citizen status in the contact line
    name_header(doc, "Riket B. Patel")
    contact_line(doc, [
        "56 Benford Ln, Edgewater Park, NJ 08010",
        "(267) 408-6295",
        "Riketpatel@gmail.com",
        "linkedin.com/in/riketpatel",
        "U.S. Citizen",
        "Trilingual: English, Hindi, Gujarati",
    ])
    headline_bar(doc, "Business Analyst  ·  Regulated Systems  ·  Public Service-Minded")

    summary_para(doc, [
        ("Started my career as a Business / Technical Analyst at Merck, was promoted through three levels (Specialist → Squad Lead → Product Manager), and spent seven years translating complex business requirements into delivered software in a heavily regulated environment. Removed ~", False),
        ("10,000 hours", True),
        (" of annual manual work for stakeholder teams across six countries and led a regulated submission program that cut processing time from ", False),
        ("6-7 months to about one month", True),
        (". M.S. Data Analytics, Penn State. New Jersey resident, U.S. citizen, NJ Real Estate Salesperson — already engaged with state systems. Drawn to this role for what the Judiciary actually does: keep records honest, court processes accessible, and the public served by systems they trust.", False),
    ])

    section_header(doc, "Selected Wins (BA Translatable)")

    bullet_with_quote(doc,
        [("Multi-system requirements coordination", True),
         (" — led the Scale IT for Clinical Portfolio program across 4 interconnected data systems (Delta Fusion, CDI, CDDR, SLS). Worked with 15+ stakeholder groups (legal-equivalent: Biostatistics, Medical Writing, Compliance, Data Stewardship) to translate business needs into delivered software. Eliminated ~10,000 hours of annual manual effort across the user base.", False)]
    )
    bullet_with_quote(doc,
        [("Multi-jurisdiction submission platform", True),
         (" — designed the A&R Submission Translation platform that cut regulated filing cycle time from ", False),
         ("6-7 months to about one month", True),
         (". User research with stakeholders in 5 countries; iterative MVP (R1.0 → R1.1). Pattern directly translates to court-record translation, eFiling localization, and accessibility-mandated multi-language workflows.", False)]
    )
    bullet_with_quote(doc,
        [("Audit-readiness SME", True),
         (" across 3 inspection-sensitive systems — documented audit trails, user-access controls, and data-integrity workflows. Same documentation discipline a court-system records integrity program requires.", False)]
    )
    bullet_with_quote(doc,
        [("BA progression", True),
         (" — promoted from entry-level Business / Technical Analyst (2018) to Specialist (2019), to Lead BA (2020-2021), to Squad Lead (2022), to Product Manager (2023). Year objectives consistently shipped on schedule with zero audit findings. 20+ peer recognitions; reverse-mentor feedback (Robert Wiley, Director, Merck):", False)],
        '"You were exceptional in how you worked with me in every area we explored. I do not have any further development feedback."'
    )
    bullet_with_quote(doc,
        [("Trilingual + community-rooted", True),
         (" — English, Hindi, Gujarati. Active member of Interfaith Organization, Next Gen Network, and capABILITY Network (disability inclusion) ERGs at Merck. Mentored Cornell Tech master's intern Jenny Mao through her 2025 summer cohort.", False)]
    )

    section_header(doc, "Experience")

    job_block(doc,
        "Business / Technical Analyst → Product Manager · Merck & Co., Inc., Upper Gwynedd, PA",
        "Aug 2018 – Nov 2025",
        [("Joined as an entry-level Business / Technical Analyst supporting BARDS clinical-data systems. Promoted through Specialist (2019), Lead BA (2020-2021), Squad Lead (2022), and Product Manager (2023). Owned requirements gathering, user research, documentation, sprint-cycle delivery, and audit-readiness across 5 interconnected systems. Subject-matter expert for ", False),
         ("regulatory audit-readiness", True),
         (" (audit trails, user access, data integrity) — the most directly transferable skill to court-records integrity work. ", False),
         ("TCO Data Steward", True),
         (" using Apptio + Ariba; Change Champion for O365 / Microsoft Teams enterprise adoption.", False)]
    )

    job_block(doc,
        "Solution Specialist II → EHR Application Specialist · NextGen Healthcare, Horsham, PA",
        "May 2014 – Apr 2017",
        [("Mirth Connect & EHR client implementations. Triaged technical issues across HL7 / FHIR integrations using ", False),
         ("SQL Server Management Studio, SQL Profiler, and SQL Tracing", True),
         (". Authored SOPs aligned to Public Health policy. Built domain expertise in ", False),
         ("HIPAA, ICD-10, Health Information Exchange, and clinical documentation", True),
         (" — same data-privacy and records-integrity disciplines a court system relies on.", False)]
    )

    job_block(doc,
        "Founder · Legacy Bridge Alliance Group (concurrent, evenings/weekends)",
        "2025 – present",
        [("Operate a small holding company alongside this job search: a New Jersey real-estate practice (HomeSmart affiliate — engaged with NJ state-licensed transactions weekly), a small-business brokerage practice, two consumer e-commerce stores, and four free public-good content sites. Available for full-time work and committed to staying local in South Jersey / Greater Philadelphia for the foreseeable future.", False)]
    )

    edu_certs_block(doc)

    footnote(doc, "References available on request. U.S. Citizen. Available to start within 2 weeks. Longer narrative at riketpatel.com.")

    out = OUTPUT_DIR / "Riket B Patel — Résumé (NJ Judiciary, BA + IT Analyst 2).docx"
    doc.save(out)
    print(f"✓ {out.name}")


# ─── BUILDER 6: NJ Judiciary cover letter ──────────────────────────────────

def build_nj_judiciary_cover_letter():
    doc = Document()
    set_margins(doc, top=0.75, bottom=0.75, left=0.85, right=0.85)

    p = add_para(doc, space_after=2, line_spacing=1.0)
    r = p.add_run("Riket B. Patel")
    set_run(r, size=19, bold=True, color=INK)
    rPr = r._r.get_or_add_rPr()
    spc = OxmlElement("w:spacing")
    spc.set(qn("w:val"), "24")
    rPr.append(spc)

    contact_line(doc, [
        "56 Benford Ln, Edgewater Park, NJ 08010",
        "(267) 408-6295",
        "Riketpatel@gmail.com",
        "U.S. Citizen",
    ])

    p = add_para(doc, space_before=16, space_after=12, line_spacing=1.3)
    r = p.add_run("May 14, 2026")
    set_run(r, size=11, color=INK)

    for line in [
        "Human Resources",
        "New Jersey Judiciary",
        "Re: Business Analyst / Information Technology Analyst 2 (Job ID 5318997)",
    ]:
        p = add_para(doc, space_after=0, line_spacing=1.3)
        r = p.add_run(line)
        set_run(r, size=11, color=INK)

    p = add_para(doc, space_before=16, space_after=8, line_spacing=1.4)
    r = p.add_run("Dear Hiring Team,")
    set_run(r, size=11.5, color=INK)

    def body_para(text_runs):
        p = add_para(doc, space_after=8, line_spacing=1.5)
        for text, bold, italic in text_runs:
            r = p.add_run(text)
            set_run(r, size=11.5, bold=bold, italic=italic, color=INK)

    def subhead(text):
        p = add_para(doc, space_before=10, space_after=4, line_spacing=1.4)
        r = p.add_run(text)
        set_run(r, size=12, bold=True, color=ACCENT)

    body_para([
        ("I am applying for the Business Analyst / Information Technology Analyst 2 role. I'd like to address upfront what your screening team will likely see first: my most recent title was Product Manager at Merck, which is more senior than this role. I want to explain why I'm applying anyway, because I'd rather be honest about that than have it raise a flag in your review.", False, False),
    ])

    subhead("Why I'm applying to this role")
    body_para([
        ("I started my career as a Business / Technical Analyst at Merck in 2018, and I spent the first four years of my Merck tenure doing exactly the work this role describes — gathering requirements from non-technical stakeholders, translating them into specifications, coordinating with delivery teams, and maintaining the documentation that compliance audits depend on. I was promoted into product management because I was good at the BA work, not because I wanted to leave it behind. Returning to a Business Analyst seat at the New Jersey Judiciary isn't a step backward for me; it's a choice to do the work that I'm best at in a place where it actually matters — for residents of the state I live in, on systems that affect real people's lives.", False, False),
    ])

    subhead("What I bring")
    body_para([
        ("Seven years inside a Fortune 100 regulated environment. I served as the subject-matter expert for ", False, False),
        ("regulatory audit-readiness", True, False),
        (" across three inspection-sensitive systems — audit trails, user access, data integrity. That's a direct skill-set match for the records-integrity and access-control work a court system requires. I led a multi-system program at Merck that eliminated about 10,000 hours of annual manual effort for stakeholder teams across six countries. I designed and shipped a multi-language submission platform that cut translated Clinical Study Reports and supporting NDA documents from about 6-7 months to about one month — the same pattern as the multi-jurisdiction, multi-language, accessibility-mandated workflows the Judiciary is modernizing.", False, False),
    ])

    body_para([
        ("On tools: I work daily in JIRA, Confluence, Mural, SQL, and the Microsoft 365 stack. I'm trilingual (English, Hindi, Gujarati) — useful in a court system where interpreter coordination and accessibility-mandated translation are real concerns.", False, False),
    ])

    subhead("Why public service, why now")
    body_para([
        ("My family came to the United States from a small village in Anand, Gujarat. We settled in Bristol, Pennsylvania, where my parents built a life from scratch. I now live in Edgewater Park, New Jersey with my wife and son. The community that raised me — teachers, courts, libraries, the people who staff the windows where you renew your driver's license — was the public sector. I want to spend a chapter of my career repaying some of that. The Judiciary's mission — an accessible, accountable court system — is one I'd genuinely be proud to point my son toward someday.", False, False),
    ])

    body_para([
        ("A practical note: I hold an active New Jersey Real Estate Salesperson license, which means I already work weekly with state-regulated systems and processes. I'm not new to the rhythms of how New Jersey government does business.", False, False),
    ])

    subhead("Commitment")
    body_para([
        ("A reasonable concern with someone whose résumé shows a senior corporate title is \"flight risk.\" I want to address it plainly: I am not interviewing this role as a stopgap. South Jersey is home; we're not relocating. The small ventures I founded after leaving Merck (a real-estate practice, a brokerage practice, and a few small consumer brands) are explicitly part-time and structured to coexist with a full-time job — I built them on nights and weekends and they were always meant to be that. If hired, I'd plan on staying long enough to ship multiple full project cycles, not 18 months.", False, False),
    ])

    body_para([
        ("My résumé, tailored to this role, is at riketpatel.com/resume/nj-judiciary/. My broader background is at riketpatel.com. References are available on request — senior leadership and peer references from Merck, all based in the United States and reachable by phone.", False, False),
    ])

    body_para([
        ("Thank you for the work the Judiciary does and for considering this application.", False, False),
    ])

    p = add_para(doc, space_before=10, space_after=4, line_spacing=1.4)
    r = p.add_run("Sincerely,")
    set_run(r, size=11.5, color=INK)

    p = add_para(doc, space_after=2, line_spacing=1.3)
    r = p.add_run("Riket B. Patel")
    set_run(r, size=11.5, bold=True, color=INK)

    for line in [
        "M.P.S. Data Analytics, Pennsylvania State University · Certified Agile Leader (CAL-1), Scrum Alliance",
        "Certified AI Product Manager · Certified Mental Health First Aider · NJ Real Estate Salesperson",
    ]:
        p = add_para(doc, space_after=0, line_spacing=1.3)
        r = p.add_run(line)
        set_run(r, size=10.5, color=MUTED)

    out = OUTPUT_DIR / "Riket B Patel — Cover Letter (NJ Judiciary, BA + IT Analyst 2).docx"
    doc.save(out)
    print(f"✓ {out.name}")


# ─── BUILDER 7: Deborah Heart & Lung Center résumé ─────────────────────────

def build_deborah_resume():
    doc = Document()
    set_margins(doc)

    name_header(doc, "Riket B. Patel")
    contact_line(doc, [
        "56 Benford Ln, Edgewater Park, NJ 08010",
        "(267) 408-6295",
        "Riketpatel@gmail.com",
        "linkedin.com/in/riketpatel",
        "U.S. Citizen",
        "Trilingual: English, Hindi, Gujarati",
    ])
    headline_bar(doc, "Leadership Development  ·  Mentorship  ·  Healthcare Operations")

    summary_para(doc, [
        ("Cross-functional leader with seven years inside a Fortune 100 regulated healthcare environment (Merck) doing the work this role asks for — informally. Promoted from Business Analyst to Squad Lead managing sprint cadence and mentorship; served as a ", False),
        ("Change Champion", True),
        (" for enterprise O365 / Teams adoption (designed and delivered training across the business); mentored a Cornell Tech master's intern to a successful summer cohort and a Director through a formal reverse-mentoring program. M.S. Data Analytics from Penn State. Certified Agile Leader (Scrum Alliance) and Mental Health First Aider. Local to Browns Mills (~30 min). Looking to professionalize what I've been doing on the side — building leaders, designing programs, and aligning learning with operational priorities — in a healthcare institution whose mission I respect.", False),
    ])

    section_header(doc, "Selected Wins (Leadership-Development Translatable)")

    bullet_with_quote(doc,
        [("Enterprise Change Champion", True),
         (" — designed and delivered training that drove organization-wide O365 and Microsoft Teams adoption at Merck. Technology roadshows, hands-on workshops, manager-level coaching sessions. Same shape as a structured behavior-change learning program.", False)]
    )
    bullet_with_quote(doc,
        [("Formal reverse-mentoring program", True),
         (" — mentored Robert Wiley, Director of Oncology National Accounts at Merck, through a multi-month structured reverse-mentoring engagement. His written assessment:", False)],
        '"You were exceptional in how you worked with me in every area we explored. I do not have any further development feedback."'
    )
    bullet_with_quote(doc,
        [("Emerging-leader mentor", True),
         (" — mentored a Cornell Tech master's student (Jenny Mao) through her 2025 Merck summer cohort. Weekly 1:1s, panel-style mock interviews, resume rewrites, career-decision coaching. Her parting note:", False)],
        '"I truly don\'t think I would have completed my internship so successfully without your mentorship."'
    )
    bullet_with_quote(doc,
        [("Squad Lead (2022) at Merck", True),
         (" — managed sprint planning, delivery cadence, and team mentorship for a cross-functional squad. Direct people-leadership experience with managers, contractors, and senior individual contributors across regulated environments. 20+ peer recognitions on Merck's INSPIRE platform — and an even higher count given to others.", False)]
    )
    bullet_with_quote(doc,
        [("ERG leadership", True),
         (" — active member of Interfaith Organization, Next Gen Network, and capABILITY Network (disability inclusion) at Merck. Mental Health First Aider certified (2024). Workforce diversity and accessibility are domains I have actually shown up in, not just discussed.", False)]
    )

    section_header(doc, "Experience")

    job_block(doc,
        "Business Analyst → Squad Lead → Product Manager · Merck & Co., Inc., Upper Gwynedd, PA",
        "Aug 2018 – Nov 2025",
        [("Joined as Business / Technical Analyst; promoted through five levels including ", False),
         ("Squad Lead", True),
         (" (2022) and Product Manager (2023). Owned cross-functional team mentorship, sprint cadence, and Change-Champion training programs across 6 countries. Subject-matter expert for regulatory audit-readiness across 3 inspection-sensitive systems — equivalent discipline to The Joint Commission and CMS audit-readiness work in a hospital setting. Aligned learning and adoption initiatives with enterprise OKR / KPI frameworks.", False)]
    )

    job_block(doc,
        "Solution Specialist II → EHR Application Specialist · NextGen Healthcare, Horsham, PA",
        "May 2014 – Apr 2017",
        [("Direct healthcare-industry experience. EHR / Mirth Connect implementations; HL7 / FHIR / SQL Server troubleshooting. Authored SOPs aligned to Public Health policy. Domain knowledge in ", False),
         ("HIPAA, ICD-10, Health Information Exchange, and clinical documentation", True),
         (" — foundation for understanding hospital operating context.", False)]
    )

    job_block(doc,
        "Founder · Legacy Bridge Alliance Group (concurrent, evenings/weekends)",
        "2025 – present",
        [("Small holding company alongside this job search: a local NJ real-estate practice (HomeSmart affiliate), a small-business brokerage practice, two consumer e-commerce stores, and four free public-good content sites (including BestPythonCourse, a free learning guide). Comfortable building free, accessible educational content in my own time — the instinct that draws me to this role.", False)]
    )

    # Custom edu/certs block — include the SHRM-SCP commitment line
    section_header(doc, "Education & Certifications")
    left = [
        [("Pennsylvania State University", True)],
        [("M.P.S. Data Analytics — May 2018", False)],
        [("Graduate Certificate, Business Analytics — Aug 2017", False)],
        [("B.S. Business Administration (Mgmt. & Marketing) — May 2014", False)],
    ]
    right = [
        [("Certified Agile Leader (CAL-1), Scrum Alliance — Jun 2024", True)],
        [("Mental Health First Aider — May 2024", True)],
        [("Certified AI Product Management — Dec 2024", True)],
        [("AWS Cloud Practitioner Path · NJ Real Estate Salesperson", False)],
        [("SHRM-SCP — committed to pursuing within Year 1", False)],
    ]
    two_column_block(doc, left, right)

    footnote(doc, "References available on request. Local to Burlington County. Available to start within 2 weeks. See cover letter for an honest note about the SHRM-SCP credential.")

    out = OUTPUT_DIR / "Riket B Patel — Résumé (Deborah Heart and Lung, Leadership Development).docx"
    doc.save(out)
    print(f"✓ {out.name}")


# ─── BUILDER 8: Deborah Heart & Lung cover letter ──────────────────────────

def build_deborah_cover_letter():
    doc = Document()
    set_margins(doc, top=0.75, bottom=0.75, left=0.85, right=0.85)

    p = add_para(doc, space_after=2, line_spacing=1.0)
    r = p.add_run("Riket B. Patel")
    set_run(r, size=19, bold=True, color=INK)
    rPr = r._r.get_or_add_rPr()
    spc = OxmlElement("w:spacing")
    spc.set(qn("w:val"), "24")
    rPr.append(spc)

    contact_line(doc, [
        "56 Benford Ln, Edgewater Park, NJ 08010",
        "(267) 408-6295",
        "Riketpatel@gmail.com",
        "U.S. Citizen",
    ])

    p = add_para(doc, space_before=16, space_after=12, line_spacing=1.3)
    r = p.add_run("May 14, 2026")
    set_run(r, size=11, color=INK)

    for line in [
        "Talent Acquisition & People Operations",
        "Deborah Heart and Lung Center",
        "Browns Mills, New Jersey",
    ]:
        p = add_para(doc, space_after=0, line_spacing=1.3)
        r = p.add_run(line)
        set_run(r, size=11, color=INK)

    p = add_para(doc, space_before=16, space_after=8, line_spacing=1.4)
    r = p.add_run("Dear Hiring Team,")
    set_run(r, size=11.5, color=INK)

    def body_para(text_runs):
        p = add_para(doc, space_after=8, line_spacing=1.5)
        for text, bold, italic in text_runs:
            r = p.add_run(text)
            set_run(r, size=11.5, bold=bold, italic=italic, color=INK)

    def subhead(text):
        p = add_para(doc, space_before=10, space_after=4, line_spacing=1.4)
        r = p.add_run(text)
        set_run(r, size=12, bold=True, color=ACCENT)

    body_para([
        ("I'm writing about the Leadership Development role at Deborah Heart and Lung Center. I want to address something important up front, because I'd rather you see it from me than have it raise a flag in screening: I do not currently hold the SHRM-SCP credential listed under \"Required\" in the posting. Everything else in the description — designing leadership programs, management training, succession planning, aligning learning with operational priorities, using engagement data to improve programs — is work I have been doing in adjacent forms for seven years inside a Fortune 100 healthcare environment. I'd like to professionalize that work, with the credential, at Deborah.", False, False),
    ])

    subhead("What I've actually done that matches this role")

    body_para([
        ("At Merck I served as the enterprise Change Champion for Microsoft 365 and Teams adoption — designed training, ran roadshows and hands-on workshops, coached managers through the change. That is a structured behavior-change learning program by another name. I served as the regulatory subject-matter expert for audit-readiness across three inspection-sensitive systems — the equivalent rigor of The Joint Commission and CMS audit-readiness work a hospital depends on. I was promoted to Squad Lead in 2022 and owned sprint cadence, delivery accountability, and mentorship for a cross-functional team that included managers and senior individual contributors.", False, False),
    ])

    body_para([
        ("On the development side specifically: I participated in Merck's formal reverse-mentoring program with Robert Wiley, Director of Oncology National Accounts. His written assessment of the engagement was ", False, False),
        ('"You were exceptional in how you worked with me in every area we explored. I do not have any further development feedback."', False, True),
        (" I mentored a Cornell Tech master's intern named Jenny Mao through her 2025 summer cohort — weekly 1:1s, panel-style mock interviews, resume rewrites, real-time coaching on hard interview questions. Her parting note: ", False, False),
        ('"I truly don\'t think I would have completed my internship so successfully without your mentorship."', False, True),
        (" Those are the conversations I am best at. They're the work this role is asking for.", False, False),
    ])

    subhead("Why Deborah specifically")

    body_para([
        ("Three reasons, plainly: I live in Edgewater Park, about thirty minutes from Browns Mills, so this is a local commitment, not a remote-arbitrage move. My only prior healthcare-industry job was at NextGen Healthcare in Horsham — I built early-career domain knowledge in HIPAA, ICD-10, and Health Information Exchange standards, so a return to the hospital side of healthcare is a return to context I already understand. And Deborah's tagline — \"healthcare is still about caring\" — matches how I personally think about service. I was raised in the Vedic tradition where work that serves others is treated as a form of practice. My father made a promise at his farewell event from our family's village in Anand, Gujarat to always give back; I am continuing that promise on my own time by setting up a Python and generative-AI teacher position there. Coming to work every day at a place whose mission centers on patient care and team-member experience is a kind of give-back I would be proud to point my son toward someday.", False, False),
    ])

    subhead("On the SHRM-SCP gap")

    body_para([
        ("I want to be specific about my commitment, not vague. Based on Merck's recent classification of my role, I believe I meet the SHRM-SCP eligibility requirements through HR-adjacent leadership work. If you hire me, I will sit for the SHRM-SCP exam within the first twelve months of employment. I would rather come into the role honest about pursuing the credential than dress up my background as conventional Human Resources when the most accurate description is \"cross-functional leader with a deep mentorship and change-management track record who is ready to make leadership development the formal job.\" If the credential is a strict gate that has to be in hand before onboarding, I understand and would withdraw the application rather than waste your screening time. If there's room for an internal commitment to earn it, I'd welcome the chance to discuss the role further.", False, False),
    ])

    subhead("What I bring on day one")

    body_para([
        ("A Master of Professional Studies in Data Analytics from Penn State — useful for the workforce-analytics piece the posting flags as preferred. Certified Agile Leader (CAL-1) from Scrum Alliance, Certified Mental Health First Aider, and a Certified AI Product Management credential from late 2024. Working knowledge of JIRA, Confluence, Mural, SQL, and the Microsoft 365 stack. Trilingual (English, Hindi, Gujarati). Active employee-resource-group experience at Merck (Interfaith Organization, Next Gen Network, and capABILITY Network for disability inclusion) — I have actually shown up in workforce-inclusion work, not just discussed it.", False, False),
    ])

    body_para([
        ("My résumé tailored to this role is at riketpatel.com/resume/deborah-heart-lung/. Five professional references from senior leadership and peer Product Management at Merck are available on request — all U.S.-based, reachable by phone, can speak directly to the mentorship and leadership-development work referenced above.", False, False),
    ])

    body_para([
        ("Thank you for considering an unconventional application. I'd welcome a conversation either way.", False, False),
    ])

    p = add_para(doc, space_before=10, space_after=4, line_spacing=1.4)
    r = p.add_run("Sincerely,")
    set_run(r, size=11.5, color=INK)

    p = add_para(doc, space_after=2, line_spacing=1.3)
    r = p.add_run("Riket B. Patel")
    set_run(r, size=11.5, bold=True, color=INK)

    for line in [
        "M.P.S. Data Analytics, Pennsylvania State University",
        "Certified Agile Leader (CAL-1), Scrum Alliance · Certified AI Product Manager",
        "Certified Mental Health First Aider · NJ Real Estate Salesperson",
    ]:
        p = add_para(doc, space_after=0, line_spacing=1.3)
        r = p.add_run(line)
        set_run(r, size=10.5, color=MUTED)

    out = OUTPUT_DIR / "Riket B Patel — Cover Letter (Deborah Heart and Lung, Leadership Development).docx"
    doc.save(out)
    print(f"✓ {out.name}")


if __name__ == "__main__":
    print(f"Building .docx files into: {OUTPUT_DIR}")
    builders = [
        ("Main résumé", build_main_resume),
        ("ADP résumé", build_adp_resume),
        ("Internet Archive résumé", build_ia_resume),
        ("Internet Archive cover letter", build_ia_cover_letter),
        ("NJ Judiciary résumé", build_nj_judiciary_resume),
        ("NJ Judiciary cover letter", build_nj_judiciary_cover_letter),
        ("Deborah Heart & Lung résumé", build_deborah_resume),
        ("Deborah Heart & Lung cover letter", build_deborah_cover_letter),
    ]
    skipped = []
    for label, fn in builders:
        try:
            fn()
        except PermissionError as e:
            print(f"⚠ SKIPPED {label} — file is open in Word. Close it and re-run to refresh.")
            skipped.append(label)
        except Exception as e:
            print(f"✗ FAILED {label}: {e}")
            skipped.append(label)
    print()
    if skipped:
        print("Done with skips:", ", ".join(skipped))
    else:
        print("Done. All files refreshed.")
